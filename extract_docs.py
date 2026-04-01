# -*- coding: utf-8 -*-
import os
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

def extract_docx_content(docx_path):
    try:
        doc = Document(docx_path)
        content = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content.append(text)
        for table in doc.tables:
            if table.rows:
                content.append("\n| " + " | ".join([cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]) + " |")
                content.append("| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |")
                for row in table.rows[1:]:
                    content.append("| " + " | ".join([cell.text.strip().replace('\n', ' ') for cell in row.cells]) + " |")
        return "\n".join(content)
    except Exception as e:
        return f"[提取失败: {str(e)}]"

base_path = r"c:\Users\陆鸣\Desktop\杆杆响天梯赛"
count = 0

for root, dirs, files in os.walk(base_path):
    for file in files:
        if file.endswith('.docx') and not file.startswith('~$'):
            docx_path = os.path.join(root, file)
            try:
                content = extract_docx_content(docx_path)
                txt_path = docx_path.replace('.docx', '_提取内容.txt')
                with open(txt_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                count += 1
                print(f"[{count}] OK: {file}")
            except Exception as e:
                print(f"[跳过] {file}")
                continue

print(f"\n完成！共提取 {count} 个文件")
